#!/usr/bin/env python3
"""
SOTE LLD Word Document Generator
Generates a professional .docx matching HLD quality.
"""
import sys
sys.stdout.reconfigure(encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────
# Colour palette (matching HLD dark-professional style in Word)
# ─────────────────────────────────────────────────────────────
C_DARK_BLUE  = RGBColor(0x1F, 0x39, 0x64)   # heading 1
C_MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # heading 2
C_LIGHT_BLUE = RGBColor(0x5B, 0x9B, 0xD5)   # heading 3
C_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
C_MID_GREY   = RGBColor(0xBF, 0xBF, 0xBF)
C_DARK_GREY  = RGBColor(0x40, 0x40, 0x40)
C_WARN_BG    = RGBColor(0xFF, 0xF2, 0xCC)   # warning box background
C_WARN_BORDER= RGBColor(0xED, 0x7D, 0x31)   # warning border
C_INFO_BG    = RGBColor(0xDA, 0xEA, 0xF3)   # info box
C_GREEN_BG   = RGBColor(0xE2, 0xEF, 0xDA)   # ok/check box
C_TBL_HEAD   = RGBColor(0x1F, 0x39, 0x64)   # table header bg
C_TBL_ALT    = RGBColor(0xDD, 0xE8, 0xF3)   # table alternating row
C_CODE_BG    = RGBColor(0x1E, 0x1E, 0x1E)   # code block background
C_CODE_FG    = RGBColor(0xD4, 0xD4, 0xD4)   # code text
C_CODE_KW    = RGBColor(0x56, 0x9C, 0xD6)   # code keyword blue

# ─────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────

def _rgb_hex(color) -> str:
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"

def set_cell_bg(cell, color: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = _rgb_hex(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in [("top", top), ("bottom", bottom),
                        ("left", left), ("right", right)]:
        if color:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), _rgb_hex(color))
            tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_border(para, color: RGBColor, left_pts=18):
    """Add left border to a paragraph (info/warning box style)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    hex_c = _rgb_hex(color)
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    str(left_pts))
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), hex_c)
    pBdr.append(left)
    pPr.append(pBdr)

def set_para_shading(para, color: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_c = _rgb_hex(color)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_c)
    pPr.append(shd)

def set_run_font(run, name="Calibri", size=10, bold=False,
                 italic=False, color=None, mono=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if mono:
        run.font.name = "Consolas"
    else:
        run.font.name = name
    if color:
        run.font.color.rgb = color

def add_heading(doc, text, level=1, numbering=None):
    """Add a heading with SOTE colour scheme."""
    p = doc.add_heading(level=level)
    p.clear()
    run = p.add_run((f"{numbering}  " if numbering else "") + text)
    sizes   = {1: 18, 2: 14, 3: 12, 4: 11}
    colors  = {1: C_DARK_BLUE, 2: C_MID_BLUE,
               3: C_LIGHT_BLUE, 4: C_DARK_GREY}
    run.bold = True
    run.font.size  = Pt(sizes.get(level, 11))
    run.font.color.rgb = colors.get(level, C_DARK_GREY)
    run.font.name  = "Calibri"
    p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text, bold=False, italic=False, size=10,
             space_after=4, indent=False):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    return p

def add_warning(doc, text, icon="⚠", bg=C_WARN_BG, border=C_WARN_BORDER):
    p   = doc.add_paragraph()
    run = p.add_run(f"{icon}  {text}")
    set_run_font(run, size=9, italic=True,
                 color=RGBColor(0x7F, 0x3F, 0x00))
    set_para_shading(p, bg)
    add_paragraph_border(p, border, left_pts=24)
    p.paragraph_format.left_indent   = Cm(0.3)
    p.paragraph_format.space_before  = Pt(4)
    p.paragraph_format.space_after   = Pt(6)
    return p

def add_note(doc, text, icon="ℹ"):
    return add_warning(doc, text, icon=icon,
                       bg=C_INFO_BG, border=C_MID_BLUE)

def add_ok(doc, text, icon="✅"):
    return add_warning(doc, text, icon=icon,
                       bg=C_GREEN_BG,
                       border=RGBColor(0x37, 0x86, 0x10))

def add_code_block(doc, code_lines):
    """Add a dark-background code block (IOS-XE style)."""
    for line in code_lines:
        p   = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name  = "Consolas"
        run.font.size  = Pt(8)
        run.font.color.rgb = C_CODE_FG
        set_para_shading(p, C_CODE_BG)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
    # spacing after block
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def add_table(doc, headers, rows, col_widths=None, alt_rows=True):
    """Add a professional table with dark header and optional alt-row shading."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=9, bold=True, color=C_WHITE)
        set_cell_bg(cell, C_TBL_HEAD)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = C_TBL_ALT if (alt_rows and ri % 2 == 0) else C_WHITE
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val) if val is not None else "")
            is_code = str(val).startswith("`") or str(val).startswith("10.") \
                      or "/" in str(val) and "." in str(val)
            set_run_font(run, size=9,
                         mono=str(val).startswith("`") or
                              any(c in str(val) for c in ["1/0/", "Hu", "Fo",
                                                          "TF", "Po", "GE",
                                                          "10.63", "10.100"]))
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

def add_page_break(doc):
    doc.add_page_break()

# ─────────────────────────────────────────────────────────────
# Cover page
# ─────────────────────────────────────────────────────────────

def build_cover(doc):
    # Top margin space
    for _ in range(4):
        doc.add_paragraph()

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SEMMELWEIS EGYETEM")
    run.font.name  = "Calibri"
    run.font.size  = Pt(14)
    run.font.color.rgb = C_MID_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Campushálózat Fejlesztési Program")
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.color.rgb = C_DARK_GREY

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Low-Level Design")
    run.font.name  = "Calibri"
    run.font.size  = Pt(28)
    run.font.color.rgb = C_DARK_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Variánsfüggetlen rész")
    run.font.name  = "Calibri"
    run.font.size  = Pt(16)
    run.font.color.rgb = C_MID_BLUE
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Enterprise Core  ·  InterConnect  ·  OOB Zóna  ·  Internet Edge + BGP")
    run.font.name  = "Calibri"
    run.font.size  = Pt(11)
    run.font.color.rgb = C_DARK_GREY
    run.italic = True

    for _ in range(3):
        doc.add_paragraph()

    # Meta table
    meta = [
        ("Verzió",     "0.1 DRAFT"),
        ("Dátum",      "2026-05-19"),
        ("Szerző",     "Uzsoki Márk"),
        ("Státusz",    "BELSŐ TERVEZET — Jóváhagyásra vár"),
        ("Forrás HLD", "SOTE-Halozatfejlesztes-HLD-V11.docx"),
        ("Hatókör",    "Variánsfüggetlen infrastruktúra rétegek"),
    ]
    t = doc.add_table(rows=len(meta), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        t.rows[i].cells[0].text = ""
        t.rows[i].cells[1].text = ""
        r1 = t.rows[i].cells[0].paragraphs[0].add_run(k)
        r2 = t.rows[i].cells[1].paragraphs[0].add_run(v)
        set_run_font(r1, size=10, bold=True, color=C_DARK_BLUE)
        set_run_font(r2, size=10, color=C_DARK_GREY)
        set_cell_bg(t.rows[i].cells[0], C_TBL_ALT)
        t.rows[i].cells[0].width = Cm(4)
        t.rows[i].cells[1].width = Cm(10)

    add_page_break(doc)

# ─────────────────────────────────────────────────────────────
# Chapter 1 — Naming conventions
# ─────────────────────────────────────────────────────────────

def build_ch1(doc):
    add_heading(doc, "Elnevezési Konvenciók", 1, "1.")

    add_heading(doc, "Hostname séma", 2, "1.1")
    add_body(doc, "Minden új eszköz neve a következő sémát követi:")
    add_code_block(doc, ["<CAMPUS>-<SZEREPKÖR>-<SORSZÁM>"])

    add_table(doc,
        ["Campus prefix", "Magyarázat"],
        [["BKT", "Belgyógyászati és Klinikai Tömb szerverterem"],
         ["KKT", "Klinikai Komplexum Tömb szerverterem"]],
        col_widths=[4, 12])

    add_table(doc,
        ["Szerepkör prefix", "Eszköztípus"],
        [["EC",      "Enterprise Core (Cisco C9500-48Y4C)"],
         ["IC",      "InterConnect switch"],
         ["IE",      "Internet Edge switch (Cisco C9200L-24P-4X)"],
         ["OOB-FW",  "OOB peremtűzfal"],
         ["OOB-TSG", "OOB Terminal Server Gateway (C1100TG-1N24P32A)"],
         ["DIST",    "Disztribúciós switch (variáns-specifikus LLD)"],
         ["ACC",     "Access switch — épület + emelet kód"]],
        col_widths=[4, 12])

    add_heading(doc, "Hostname példák", 3, "1.1.1")
    add_table(doc,
        ["Hostname", "Eszköz / Szerepkör"],
        [["BKT-EC-1",      "BKT Enterprise Core #1"],
         ["BKT-EC-2",      "BKT Enterprise Core #2"],
         ["KKT-EC-1",      "KKT Enterprise Core #1"],
         ["KKT-EC-2",      "KKT Enterprise Core #2"],
         ["BKT-IC-1",      "BKT InterConnect #1"],
         ["BKT-IC-2",      "BKT InterConnect #2"],
         ["KKT-IC-1",      "KKT InterConnect #1"],
         ["KKT-IC-2",      "KKT InterConnect #2"],
         ["BKT-IE-1",      "BKT Internet Edge switch"],
         ["KKT-IE-1",      "KKT Internet Edge switch"],
         ["BKT-OOB-FW",    "BKT OOB peremtűzfal"],
         ["KKT-OOB-FW",    "KKT OOB peremtűzfal"],
         ["BKT-OOB-TSG",   "BKT Terminal Server Gateway"],
         ["KKT-OOB-TSG",   "KKT Terminal Server Gateway"]],
        col_widths=[5, 11])

    add_note(doc,
             "Disztribúciós eszközök: <ÉPÜLET>-DIST-1 / <ÉPÜLET>-DIST-2 "
             "— pl. HOGYES-DIST-1, BALASSA-DIST-2. "
             "Pontos névlista a variáns-specifikus LLD-ben.")

    add_heading(doc, "Interface Description séma", 2, "1.2")
    add_body(doc, "Minden fizikai és logikai interface description mezőjének formátuma:")
    add_code_block(doc, ["TO-<CÉLESZKÖZ-HOSTNAME>-<CÉLPORT>  [ <FUNKCIÓ> ]"])
    add_table(doc,
        ["Példa", "Leírás"],
        [["TO-BKT-IC-1-Fo1/0/1 [BACKBONE-40G]",     "BKT-EC-1 portján, BKT-IC-1 felé"],
         ["TO-BKT-OOB-FW-Gi1/0/1 [OOB-MGMT]",       "OOB tűzfal felé"],
         ["TO-PROMC-RTR-BKT [BGP-PRIMARY]",           "SP/Pro M router felé (Internet Edge)"],
         ["TO-BKATA-Gi6/1 [KOEGZISZTENCIA-MIGRACIO]","Legacy core felé, migráció idején"]],
        col_widths=[8, 8])

    add_heading(doc, "VLAN névséma", 2, "1.3")
    add_body(doc,
             "Meglévő VLAN-ok neve változatlan marad. "
             "Új VLAN-ok sémája: SOTE-<CSOPORT>-<FUNKCIÓ>")
    add_table(doc,
        ["VLAN ID", "Név", "Funkció"],
        [["2000", "SOTE-DOLG-CAMPUS",           "Általános dolgozók"],
         ["2001", "SOTE-MEDI-CAMPUS",           "Orvosi/klinikai dolgozók"],
         ["2002", "SOTE-TANU-CAMPUS",           "Hallgatók"],
         ["2003", "SOTE-GAZD-CAMPUS",           "Gazdasági dolgozók"],
         ["2004", "SOTE-MEDI-TANU-CAMPUS",      "Kombinált MEDI+TANU"],
         ["2005", "SOTE-MEDI-GAZD-CAMPUS",      "Kombinált MEDI+GAZD"],
         ["2006", "SOTE-MEDI-GAZD-TANU-CAMPUS", "Hármas kombináció"],
         ["2007", "SOTE-GAZD-TANU-CAMPUS",      "Kombinált GAZD+TANU"],
         ["2008", "SOTE-INFI-CAMPUS",           "IT munkatársak"],
         ["2009", "SOTE-BIZT-CAMPUS",           "Biztonságtechnika"],
         ["422",  "SOTE-INFADMIN-MGMT",         "IT adminisztrátorok"],
         ["220",  "SOTE-MEDMOB-IOT",            "Mobil medikai eszközök"],
         ["900",  "SOTE-OOB-MGMT",              "Out-of-Band management"],
         ["901",  "SOTE-INFRA-P2P",             "Infrastructure P2P linkek"]],
        col_widths=[2.5, 7, 7])

    add_heading(doc, "ACL / Prefix-list / Route-map névséma", 2, "1.4")
    add_code_block(doc, ["<ESZKÖZ>-<IRÁNY>-<CÉL>-<TÍPUS>"])
    add_table(doc,
        ["Példa", "Leírás"],
        [["BKT-EC-1-OUT-CAMPUS-PL",   "BKT-EC-1-ről kifelé, campus, prefix-list"],
         ["SOTE-BGP-OUT-PROMC-RM",    "BGP kifelé Pro M felé, route-map"],
         ["SOTE-BGP-IN-PROMC-RM",     "BGP befelé Pro M felől, route-map"],
         ["SOTE-OSPF100-REDIST-RM",   "OSPF 100 redistributálási route-map"]],
        col_widths=[7, 9])

# ─────────────────────────────────────────────────────────────
# Chapter 2 — IP Address Plan
# ─────────────────────────────────────────────────────────────

def build_ch2(doc):
    add_heading(doc, "IP-Cím Terv", 1, "2.")

    add_heading(doc, "Meglévő tartományok — érintetlen", 2, "2.1")
    add_table(doc,
        ["Tartomány", "Szerepkör", "Megjegyzés"],
        [["10.63.64.0/21",  "Management (meglévő)",      "Nexus mgmt VRF, default GW: 10.63.64.1"],
         ["172.27.0.0/24",  "VXLAN underlay loopbackok", "Nexus fabric, OSPF Process 2"],
         ["172.27.4.0/24",  "VXLAN P2P linkek",          "Nexus fabric underlay"],
         ["193.6.209.0/24", "BGP / Internet peering",    "AS 65008, nyilvános"],
         ["10.10.255.0/30", "OSPF Process 1 (Forti–BKATA)", "Megmarad migráció végéig"]],
        col_widths=[4, 5, 7.5])

    add_heading(doc, "Új infrastruktúra tartományok", 2, "2.2")

    add_heading(doc, "Management — új eszközök", 3, "2.2.1")
    add_table(doc,
        ["Tartomány", "Eszközök", "Prefix"],
        [["BKT új infrastruktúra mgmt",
          "BKT-EC-1/2, BKT-IC-1/2, BKT-IE-1, BKT-OOB-FW/TSG",
          "10.63.72.0/24"],
         ["KKT új infrastruktúra mgmt",
          "KKT-EC-1/2, KKT-IC-1/2, KKT-IE-1, KKT-OOB-FW/TSG",
          "10.63.73.0/24"],
         ["OOB dedikált hálózat",
          "OOB-FW ↔ OOB-TSG, admin VPN hozzáférés",
          "10.63.74.0/24"]],
        col_widths=[5, 8, 3.5])

    add_heading(doc, "BKT Management IP kiosztás (10.63.72.0/24)", 3, "2.2.2")
    add_table(doc,
        ["IP", "Hostname", "Szerepkör"],
        [["10.63.72.1",  "—",           "Default gateway (Management VLAN SVI)"],
         ["10.63.72.10", "BKT-EC-1",    "Enterprise Core #1 Management"],
         ["10.63.72.11", "BKT-EC-2",    "Enterprise Core #2 Management"],
         ["10.63.72.14", "BKT-IC-1",    "InterConnect #1 Management"],
         ["10.63.72.15", "BKT-IC-2",    "InterConnect #2 Management"],
         ["10.63.72.18", "BKT-IE-1",    "Internet Edge switch Management"],
         ["10.63.72.20", "BKT-OOB-FW",  "OOB peremtűzfal Management"],
         ["10.63.72.22", "BKT-OOB-TSG", "Terminal Server Gateway Management"],
         ["10.63.72.200–254", "—",      "DHCP / dinamikus (fenntartott)"]],
        col_widths=[4, 4, 8.5])

    add_heading(doc, "KKT Management IP kiosztás (10.63.73.0/24)", 3, "2.2.3")
    add_table(doc,
        ["IP", "Hostname", "Szerepkör"],
        [["10.63.73.1",  "—",           "Default gateway (Management VLAN SVI)"],
         ["10.63.73.10", "KKT-EC-1",    "Enterprise Core #1 Management"],
         ["10.63.73.11", "KKT-EC-2",    "Enterprise Core #2 Management"],
         ["10.63.73.14", "KKT-IC-1",    "InterConnect #1 Management"],
         ["10.63.73.15", "KKT-IC-2",    "InterConnect #2 Management"],
         ["10.63.73.18", "KKT-IE-1",    "Internet Edge switch Management"],
         ["10.63.73.20", "KKT-OOB-FW",  "OOB peremtűzfal Management"],
         ["10.63.73.22", "KKT-OOB-TSG", "Terminal Server Gateway Management"]],
        col_widths=[4, 4, 8.5])

    add_heading(doc, "Infrastructure Loopbackok (Router-ID-ok)", 3, "2.2.4")
    add_table(doc,
        ["IP", "Hostname", "Szerepkör"],
        [["10.63.75.1/32", "BKT-EC-1", "OSPF Router-ID / Loopback0"],
         ["10.63.75.2/32", "BKT-EC-2", "OSPF Router-ID / Loopback0"],
         ["10.63.75.3/32", "KKT-EC-1", "OSPF Router-ID / Loopback0"],
         ["10.63.75.4/32", "KKT-EC-2", "OSPF Router-ID / Loopback0"],
         ["10.63.75.5/32", "BKT-IC-1", "OSPF Router-ID / Loopback0"],
         ["10.63.75.6/32", "BKT-IC-2", "OSPF Router-ID / Loopback0"],
         ["10.63.75.7/32", "KKT-IC-1", "OSPF Router-ID / Loopback0"],
         ["10.63.75.8/32", "KKT-IC-2", "OSPF Router-ID / Loopback0"]],
        col_widths=[4, 4, 8.5])

    add_heading(doc, "Infrastructure P2P linkek (10.63.76.0/24)", 3, "2.2.5")
    add_table(doc,
        ["Subnet /30", "A-oldal (.1)", "B-oldal (.2)", "Funkció"],
        [["10.63.76.0/30",  "BKT-EC-1", "BKT-EC-2", "EC#1↔EC#2 BKT belső gerinc"],
         ["10.63.76.4/30",  "KKT-EC-1", "KKT-EC-2", "EC#1↔EC#2 KKT belső gerinc"],
         ["10.63.76.8/30",  "BKT-EC-1", "BKT-IC-1", "BKT EC#1 → IC#1"],
         ["10.63.76.12/30", "BKT-EC-1", "BKT-IC-2", "BKT EC#1 → IC#2"],
         ["10.63.76.16/30", "BKT-EC-2", "BKT-IC-1", "BKT EC#2 → IC#1"],
         ["10.63.76.20/30", "BKT-EC-2", "BKT-IC-2", "BKT EC#2 → IC#2"],
         ["10.63.76.24/30", "KKT-EC-1", "KKT-IC-1", "KKT EC#1 → IC#1"],
         ["10.63.76.28/30", "KKT-EC-1", "KKT-IC-2", "KKT EC#1 → IC#2"],
         ["10.63.76.32/30", "KKT-EC-2", "KKT-IC-1", "KKT EC#2 → IC#1"],
         ["10.63.76.36/30", "KKT-EC-2", "KKT-IC-2", "KKT EC#2 → IC#2"],
         ["10.63.76.40/30", "BKT-IC-1", "KKT-IC-1", "IC cross-site BKT#1↔KKT#1"],
         ["10.63.76.44/30", "BKT-IC-1", "KKT-IC-2", "IC cross-site BKT#1↔KKT#2"],
         ["10.63.76.48/30", "BKT-IC-2", "KKT-IC-1", "IC cross-site BKT#2↔KKT#1"],
         ["10.63.76.52/30", "BKT-IC-2", "KKT-IC-2", "IC cross-site BKT#2↔KKT#2"],
         ["10.63.76.56/30", "BKT-OOB-FW", "BKT-IC-1", "OOB-FW BKT → IC#1"],
         ["10.63.76.60/30", "BKT-OOB-FW", "BKT-IC-2", "OOB-FW BKT → IC#2"],
         ["10.63.76.64/30", "KKT-OOB-FW", "KKT-IC-1", "OOB-FW KKT → IC#1"],
         ["10.63.76.68/30", "KKT-OOB-FW", "KKT-IC-2", "OOB-FW KKT → IC#2"],
         ["10.63.76.100/30","BKT-EC-1",   "Fortigate BKT", "Zónahatár Po2 BKT-EC-1"],
         ["10.63.76.104/30","BKT-EC-2",   "Fortigate BKT", "Zónahatár Po2 BKT-EC-2"],
         ["10.63.76.108/30","KKT-EC-1",   "Fortigate KKT", "Zónahatár Po2 KKT-EC-1"],
         ["10.63.76.112/30","KKT-EC-2",   "Fortigate KKT", "Zónahatár Po2 KKT-EC-2"]],
        col_widths=[3.5, 3.5, 3.5, 6])

    add_note(doc,
             "Disztribúciós uplink P2P linkek (D1 és D2) külön /24 allokáció — "
             "a variáns-specifikus LLD-ben kerül rögzítésre.")

    add_heading(doc, "Kliens Campus VLAN-ok — ISE DHCP scope-ok", 3, "2.2.6")
    add_table(doc,
        ["VLAN ID", "Név", "Subnet", "Hosts", "Gateway"],
        [["2000", "SOTE-DOLG-CAMPUS",           "10.100.0.0/22",  "1022", "10.100.0.1"],
         ["2001", "SOTE-MEDI-CAMPUS",           "10.100.4.0/22",  "1022", "10.100.4.1"],
         ["2002", "SOTE-TANU-CAMPUS",           "10.100.8.0/22",  "1022", "10.100.8.1"],
         ["2003", "SOTE-GAZD-CAMPUS",           "10.100.12.0/22", "1022", "10.100.12.1"],
         ["2004", "SOTE-MEDI-TANU-CAMPUS",      "10.100.16.0/23", "510",  "10.100.16.1"],
         ["2005", "SOTE-MEDI-GAZD-CAMPUS",      "10.100.18.0/23", "510",  "10.100.18.1"],
         ["2006", "SOTE-MEDI-GAZD-TANU-CAMPUS", "10.100.20.0/23", "510",  "10.100.20.1"],
         ["2007", "SOTE-GAZD-TANU-CAMPUS",      "10.100.22.0/23", "510",  "10.100.22.1"],
         ["2008", "SOTE-INFI-CAMPUS",           "10.100.24.0/24", "254",  "10.100.24.1"],
         ["2009", "SOTE-BIZT-CAMPUS",           "10.100.25.0/24", "254",  "10.100.25.1"],
         ["422",  "SOTE-INFADMIN-MGMT",         "10.100.26.0/24", "254",  "10.100.26.1"],
         ["220",  "SOTE-MEDMOB-IOT",            "10.100.27.0/24", "254",  "10.100.27.1"]],
        col_widths=[1.8, 6, 3.5, 1.8, 3.5])

    add_heading(doc, "OSPF terv — új campus infrastruktúra", 3, "2.2.7")
    add_table(doc,
        ["Paraméter", "Érték", "Indok"],
        [["Process ID",     "100",                "Elkülönített — nem ütközik meglévő OSPF 1/2-vel"],
         ["Area",           "0.0.0.0 (backbone)", "HLD hibájának elkerülése — meglévő OSPF 1/2 nem backbone area!"],
         ["Router-ID",      "Loopback0 (10.63.75.x)", "Stabil, eszközönként egyedi"],
         ["Network type",   "point-to-point",     "P2P /30 linken — nincs DR/BDR overhead"],
         ["Hello / Dead",   "10s / 40s",          "Campus-szintű konvergencia (default IOS-XE érték)"],
         ["Authentication", "MD5 (SOTE-OSPF-KEY)","Hitelesített szomszédság"],
         ["Redistribute",   "connected + static → OSPF 100", "Campus útvonalak hirdetése"]],
        col_widths=[4, 5, 7.5])

    add_warning(doc,
                "A meglévő OSPF Process 1 (Fortigate–BKATA, Area 10.10.255.1) és "
                "OSPF Process 2 (VXLAN underlay, Area 172.27.0.0) mindkettő "
                "non-backbone area — ez az új OSPF 100 bevezetésének egyik fő indoka. "
                "Az OSPF Process 1 a Fortigate–új EC integrációig megmarad.")

# ─────────────────────────────────────────────────────────────
# Chapter 3 — Enterprise Core
# ─────────────────────────────────────────────────────────────

def build_ch3(doc):
    add_heading(doc, "Enterprise Core", 1, "3.")

    add_heading(doc, "Eszköz áttekintés", 2, "3.1")
    add_table(doc,
        ["Hostname", "Modell", "Helyszín", "Szerepkör"],
        [["BKT-EC-1", "Cisco Catalyst C9500-48Y4C", "BKT Szerverterem", "Enterprise Core #1"],
         ["BKT-EC-2", "Cisco Catalyst C9500-48Y4C", "BKT Szerverterem", "Enterprise Core #2"],
         ["KKT-EC-1", "Cisco Catalyst C9500-48Y4C", "KKT Szerverterem", "Enterprise Core #1"],
         ["KKT-EC-2", "Cisco Catalyst C9500-48Y4C", "KKT Szerverterem", "Enterprise Core #2"]],
        col_widths=[3.5, 5.5, 4.5, 4])

    add_heading(doc, "C9500-48Y4C portok", 3, "3.1.1")
    add_table(doc,
        ["Port jelölés", "Típus", "Darab", "Alkalmazás"],
        [["TwentyFiveGigE1/0/1–48", "25G SFP28", "48",
          "Fortigate (MPO breakout), Disztribúció D1/D2 uplinkok, OOB-TSG"],
         ["HundredGigE1/0/49–52",   "100G QSFP28", "4",
          "EC↔EC DAC (100G), IC uplinkok (40G QSFP-40G-SR4)"],
         ["GigabitEthernet0",        "1G RJ45", "1",
          "Dedikált OOB Management port (Mgmt-vrf)"]],
        col_widths=[5, 3, 1.5, 7])

    add_heading(doc, "Fizikai port-kiosztás", 2, "3.2")

    add_heading(doc, "BKT-EC-1 port-kiosztás", 3, "3.2.1")
    add_table(doc,
        ["Port", "Típus", "Céleszköz", "Cél-port", "Médium", "SFP/Kábel"],
        [["Hu1/0/49",  "100G QSFP28", "BKT-EC-2",    "Hu1/0/49",  "100G DAC", "QSFP28-100G-CU"],
         ["Hu1/0/50",  "100G QSFP28", "BKT-EC-2",    "Hu1/0/50",  "100G DAC", "QSFP28-100G-CU"],
         ["Fo1/0/51",  "40G QSFP+",  "BKT-IC-1",    "Fo1/0/1",   "40G MMF",  "QSFP-40G-SR4"],
         ["Fo1/0/52",  "40G QSFP+",  "BKT-IC-2",    "Fo1/0/1",   "40G MMF",  "QSFP-40G-SR4"],
         ["TF1/0/1",   "10G SFP28",  "Fortigate BKT","—",         "10G MMF (MPO lane 1)", "SFP-10G-SR"],
         ["TF1/0/2",   "10G SFP28",  "Fortigate BKT","—",         "10G MMF (MPO lane 2)", "SFP-10G-SR"],
         ["TF1/0/3–9", "10G SFP28",  "KKT Dist D2 #1–7","D2-uplink","10G SMF", "SFP-10G-LR"],
         ["TF1/0/10–20","10G SFP28", "BKT Dist D1 #1–11","D1-uplink","10G SMF","SFP-10G-LR"],
         ["TF1/0/47",  "10G SFP28",  "BKT-OOB-TSG", "Gi1/0/1",   "10G MMF",  "SFP-10G-SR"],
         ["GE0",       "1G Mgmt",    "OOB switch",  "—",          "UTP Cat6", "—"]],
        col_widths=[2.5, 2.5, 3.5, 2.5, 3, 3])

    add_heading(doc, "BKT-EC-2 port-kiosztás", 3, "3.2.2")
    add_table(doc,
        ["Port", "Típus", "Céleszköz", "Cél-port", "Médium", "SFP/Kábel"],
        [["Hu1/0/49",  "100G QSFP28", "BKT-EC-1",    "Hu1/0/49",  "100G DAC", "QSFP28-100G-CU"],
         ["Hu1/0/50",  "100G QSFP28", "BKT-EC-1",    "Hu1/0/50",  "100G DAC", "QSFP28-100G-CU"],
         ["Fo1/0/51",  "40G QSFP+",  "BKT-IC-1",    "Fo1/0/2",   "40G MMF",  "QSFP-40G-SR4"],
         ["Fo1/0/52",  "40G QSFP+",  "BKT-IC-2",    "Fo1/0/2",   "40G MMF",  "QSFP-40G-SR4"],
         ["TF1/0/1",   "10G SFP28",  "Fortigate BKT","—",         "10G MMF (MPO lane 3)", "SFP-10G-SR"],
         ["TF1/0/2",   "10G SFP28",  "Fortigate BKT","—",         "10G MMF (MPO lane 4)", "SFP-10G-SR"],
         ["TF1/0/3–9", "10G SFP28",  "KKT Dist D2 #8–14","D2-uplink","10G SMF","SFP-10G-LR"],
         ["TF1/0/10–20","10G SFP28", "BKT Dist D1 #1–11","D1-uplink","10G SMF","SFP-10G-LR"],
         ["TF1/0/47",  "10G SFP28",  "BKT-OOB-TSG", "Gi1/0/2",   "10G MMF",  "SFP-10G-SR"],
         ["GE0",       "1G Mgmt",    "OOB switch",  "—",          "UTP Cat6", "—"]],
        col_widths=[2.5, 2.5, 3.5, 2.5, 3, 3])

    add_note(doc,
             "KKT-EC-1 és KKT-EC-2: azonos logika, tükrözve. "
             "KKT-EC-1: BKT Dist D2 #1–6 + KKT Dist D1 #1–14 (helyi). "
             "KKT-EC-2: BKT Dist D2 #7–11 + KKT Dist D1 #1–14 (helyi). "
             "Pontos épület–port mapping az optikai felmérés után, variáns-specifikus LLD-ben.")

    add_warning(doc,
                "A Fortigate → EC kapcsolat MPO-LC breakout kábelen fut: "
                "a Fortigate QSFP-40G-SR4 (MPO-12) portja 2× LC duplex párba "
                "bomlik szét, amelyek a EC TF1/0/1–2 portjaira csatlakoznak "
                "(2×10G = Po2 LACP LAG). Ha a Fortigate ennél nagyobb "
                "sávszélességet igényel, 4× lane (4×10G = Po2 40G) is kialakítható.")

    add_heading(doc, "IOS-XE konfiguráció — BKT-EC-1", 2, "3.3")
    add_note(doc,
             "Sablon: BKT-EC-2, KKT-EC-1, KKT-EC-2 ugyanígy — "
             "IP-értékek az IP-cím tervnek megfelelően cserélendők.")

    add_heading(doc, "Alap konfiguráció", 3, "3.3.1")
    add_code_block(doc, [
        "hostname BKT-EC-1",
        "!",
        "ip domain-name sote.hu",
        "ip name-server 10.63.64.2",
        "!",
        "! ── Authentikáció ──────────────────────────────────",
        "aaa new-model",
        "aaa authentication login default group tacacs+ local",
        "aaa authorization exec default group tacacs+ local",
        "aaa accounting exec default start-stop group tacacs+",
        "!",
        "username admin privilege 15 algorithm-type sha256 secret <PASSWORD>",
        "!",
        "! ── SSH ────────────────────────────────────────────",
        "crypto key generate rsa modulus 4096",
        "ip ssh version 2",
        "ip ssh time-out 60",
        "ip ssh authentication-retries 3",
        "line vty 0 15",
        " transport input ssh",
        " login authentication default",
        " exec-timeout 15 0",
        "!",
        "! ── NTP ────────────────────────────────────────────",
        "ntp server 10.63.64.5 prefer",
        "ntp server 10.63.64.6",
        "!",
        "! ── Logging ────────────────────────────────────────",
        "logging host 10.63.64.10",
        "logging trap informational",
        "logging source-interface Loopback0",
        "service timestamps log datetime msec localtime",
        "!",
        "! ── SNMP ───────────────────────────────────────────",
        "snmp-server community <COMMUNITY-RO> RO",
        "snmp-server location BKT Szerverterem",
        "snmp-server contact netops@sote.hu",
        "snmp-server host 10.63.64.10 version 2c <COMMUNITY-RO>",
    ])

    add_heading(doc, "Management interface (OOB)", 3, "3.3.2")
    add_code_block(doc, [
        "interface GigabitEthernet0",
        " description TO-BKT-OOB-TSG-Gi0/0 [OOB-MGMT]",
        " vrf forwarding Mgmt-vrf",
        " ip address 10.63.72.10 255.255.255.0",
        " no shutdown",
        "!",
        "ip route vrf Mgmt-vrf 0.0.0.0 0.0.0.0 10.63.72.1",
    ])

    add_heading(doc, "Loopback0 (Router-ID)", 3, "3.3.3")
    add_code_block(doc, [
        "interface Loopback0",
        " description ROUTER-ID-OSPF100",
        " ip address 10.63.75.1 255.255.255.255",
        " no shutdown",
    ])

    add_heading(doc, "EC#1 ↔ EC#2 belső gerinc (Port-Channel1, 2×100G DAC)", 3, "3.3.4")
    add_code_block(doc, [
        "interface Port-channel1",
        " description TO-BKT-EC-2-Po1 [BACKBONE-100G]",
        " no switchport",
        " ip address 10.63.76.1 255.255.255.252",
        " ip ospf network point-to-point",
        " ip ospf 100 area 0",
        " ip ospf authentication message-digest",
        " ip ospf message-digest-key 1 md5 <OSPF-KEY>",
        " no shutdown",
        "!",
        "interface HundredGigE1/0/49",
        " description TO-BKT-EC-2-Hu1/0/49 [BACKBONE-DAC-1]",
        " no switchport",
        " channel-group 1 mode active",
        " no shutdown",
        "!",
        "interface HundredGigE1/0/50",
        " description TO-BKT-EC-2-Hu1/0/50 [BACKBONE-DAC-2]",
        " no switchport",
        " channel-group 1 mode active",
        " no shutdown",
    ])

    add_heading(doc, "EC → InterConnect uplinkok (40G MMF)", 3, "3.3.5")
    add_code_block(doc, [
        "interface FortyGigabitEthernet1/0/51",
        " description TO-BKT-IC-1-Fo1/0/1 [BACKBONE-40G]",
        " no switchport",
        " ip address 10.63.76.9 255.255.255.252",
        " ip ospf network point-to-point",
        " ip ospf 100 area 0",
        " ip ospf authentication message-digest",
        " ip ospf message-digest-key 1 md5 <OSPF-KEY>",
        " no shutdown",
        "!",
        "interface FortyGigabitEthernet1/0/52",
        " description TO-BKT-IC-2-Fo1/0/1 [BACKBONE-40G]",
        " no switchport",
        " ip address 10.63.76.13 255.255.255.252",
        " ip ospf network point-to-point",
        " ip ospf 100 area 0",
        " ip ospf authentication message-digest",
        " ip ospf message-digest-key 1 md5 <OSPF-KEY>",
        " no shutdown",
    ])

    add_heading(doc, "EC → Fortigate zónahatár (Po2, 2×10G MPO breakout)", 3, "3.3.6")
    add_code_block(doc, [
        "interface Port-channel2",
        " description TO-FORTIGATE-BKT-Po2 [ZONEBORDER-FW]",
        " no switchport",
        " ip address 10.63.76.101 255.255.255.252",
        " no shutdown",
        "!",
        "interface TwentyFiveGigE1/0/1",
        " description TO-FORTIGATE-BKT-MPO-LANE1 [ZONEBORDER-FW]",
        " no switchport",
        " channel-group 2 mode active",
        " no shutdown",
        "!",
        "interface TwentyFiveGigE1/0/2",
        " description TO-FORTIGATE-BKT-MPO-LANE2 [ZONEBORDER-FW]",
        " no switchport",
        " channel-group 2 mode active",
        " no shutdown",
    ])
    add_note(doc,
             "OSPF Process 1 (Fortigate–EC szomszédság) a Fortigate meglévő OSPF 1 "
             "konfigurációjával kompatibilis módon kerül kialakításra. "
             "A BKATA–Fortigate OSPF 1 szomszédság a migráció végéig párhuzamosan fut.")

    add_heading(doc, "OSPF Process 100", 3, "3.3.7")
    add_code_block(doc, [
        "router ospf 100",
        " router-id 10.63.75.1",
        " area 0 authentication message-digest",
        " passive-interface default",
        " no passive-interface Port-channel1",
        " no passive-interface FortyGigabitEthernet1/0/51",
        " no passive-interface FortyGigabitEthernet1/0/52",
        " no passive-interface TwentyFiveGigE1/0/3",
        " !  ... (disztribúciós uplink interfészek — ismételni minden aktív portnál)",
        " network 10.63.75.1 0.0.0.0 area 0",
        " network 10.63.76.0 0.0.0.255 area 0",
        " redistribute connected subnets route-map OSPF100-REDIST-CONNECTED",
        "!",
        "ip prefix-list OSPF100-LOOPBACKS seq 10 permit 10.63.75.0/24 le 32",
        "ip prefix-list OSPF100-INFRA    seq 10 permit 10.63.76.0/24 le 30",
        "!",
        "route-map OSPF100-REDIST-CONNECTED permit 10",
        " match ip address prefix-list OSPF100-LOOPBACKS OSPF100-INFRA",
    ])

    add_heading(doc, "Enterprise Core — összefoglaló IP-táblázat", 2, "3.4")
    add_table(doc,
        ["Eszköz", "Interface", "IP", "Leírás"],
        [["BKT-EC-1", "Loopback0",          "10.63.75.1/32",   "OSPF Router-ID"],
         ["BKT-EC-1", "GE0 (Mgmt-vrf)",     "10.63.72.10/24",  "OOB Management"],
         ["BKT-EC-1", "Po1 (↔BKT-EC-2)",    "10.63.76.1/30",   "EC belső gerinc"],
         ["BKT-EC-1", "Fo1/0/51 (↔IC#1)",   "10.63.76.9/30",   "IC#1 uplink"],
         ["BKT-EC-1", "Fo1/0/52 (↔IC#2)",   "10.63.76.13/30",  "IC#2 uplink"],
         ["BKT-EC-1", "Po2 (↔Fortigate)",   "10.63.76.101/30", "Zónahatár Forti"],
         ["BKT-EC-2", "Loopback0",          "10.63.75.2/32",   "OSPF Router-ID"],
         ["BKT-EC-2", "GE0 (Mgmt-vrf)",     "10.63.72.11/24",  "OOB Management"],
         ["BKT-EC-2", "Po1 (↔BKT-EC-1)",    "10.63.76.2/30",   "EC belső gerinc"],
         ["BKT-EC-2", "Fo1/0/51 (↔IC#1)",   "10.63.76.17/30",  "IC#1 uplink"],
         ["BKT-EC-2", "Fo1/0/52 (↔IC#2)",   "10.63.76.21/30",  "IC#2 uplink"],
         ["BKT-EC-2", "Po2 (↔Fortigate)",   "10.63.76.105/30", "Zónahatár Forti"],
         ["KKT-EC-1", "Loopback0",          "10.63.75.3/32",   "OSPF Router-ID"],
         ["KKT-EC-1", "GE0 (Mgmt-vrf)",     "10.63.73.10/24",  "OOB Management"],
         ["KKT-EC-1", "Po1 (↔KKT-EC-2)",    "10.63.76.5/30",   "EC belső gerinc"],
         ["KKT-EC-1", "Fo1/0/51 (↔IC#1)",   "10.63.76.25/30",  "IC#1 uplink"],
         ["KKT-EC-1", "Fo1/0/52 (↔IC#2)",   "10.63.76.29/30",  "IC#2 uplink"],
         ["KKT-EC-1", "Po2 (↔Fortigate)",   "10.63.76.109/30", "Zónahatár Forti"],
         ["KKT-EC-2", "Loopback0",          "10.63.75.4/32",   "OSPF Router-ID"],
         ["KKT-EC-2", "GE0 (Mgmt-vrf)",     "10.63.73.11/24",  "OOB Management"],
         ["KKT-EC-2", "Po1 (↔KKT-EC-1)",    "10.63.76.6/30",   "EC belső gerinc"],
         ["KKT-EC-2", "Fo1/0/51 (↔IC#1)",   "10.63.76.33/30",  "IC#1 uplink"],
         ["KKT-EC-2", "Fo1/0/52 (↔IC#2)",   "10.63.76.37/30",  "IC#2 uplink"],
         ["KKT-EC-2", "Po2 (↔Fortigate)",   "10.63.76.113/30", "Zónahatár Forti"]],
        col_widths=[3, 4, 3.5, 6])

# ─────────────────────────────────────────────────────────────
# Build & Save
# ─────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    build_cover(doc)
    build_ch1(doc)
    add_page_break(doc)
    build_ch2(doc)
    add_page_break(doc)
    build_ch3(doc)

    out = r"c:\Projects\AI\Prompting\SOTE-LLD-V01.docx"
    doc.save(out)
    print(f"[OK] Saved: {out}")

if __name__ == "__main__":
    main()
